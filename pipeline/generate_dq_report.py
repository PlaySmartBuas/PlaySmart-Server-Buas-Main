import argparse
import json
import re
from pathlib import Path
from datetime import datetime
from statistics import mean

import pandas as pd
from docx import Document

# --------------------------
# Filename parsing
# --------------------------

# Supports:
#   - 2nd_game_P004_..._dq_metrics.json          (old)
#   - 2nd_game_Pc3_V_P004_..._dq_metrics.json    (new, with PC info)
FILENAME_RE = re.compile(
    r'(?P<match_prefix>\d+(?:st|nd|rd|th)_game)_(?:Pc(?P<pc>\d+)_V_)?(?P<player>P\d{3})_.*?_dq_metrics\.json$',
    re.IGNORECASE,
)


def parse_ids_from_filename(path: Path):
    """
    Extract match prefix (e.g., '6th_game'), player id (e.g., 'P044'),
    and optional PC id (e.g., '3') from filename.

    Returns (match_prefix, player, pc) or (None, None, None) if not matched.
    """
    m = FILENAME_RE.search(path.name)
    if not m:
        return None, None, None
    return m.group("match_prefix"), m.group("player"), m.group("pc")


# --------------------------
# KPI helpers & modality mapping
# --------------------------

# Column groups for modality-level completeness,
# based on the merged __cleaned.csv structure you provided.
GAZE_COLS = {
    "left_gaze_x",
    "left_gaze_y",
    "right_gaze_x",
    "right_gaze_y",
    "screen_x_norm",
    "screen_y_norm",
}

INPUT_COLS = {
    "event_type",
    "details_x",
    "details_y",
    "duration",
    "screen_x",
    "screen_y",
}

EMOTION_COLS = {
    "emotion",
    "confidence",
}

# Metadata / derived columns we do NOT want in completeness calculations
META_COLS = {
    "unix_time",
    "time_offset_ms",
    "datetime",
    "datetime_parsed",
}


def pct(x):
    """Return x rounded to 2 decimals if not None/NaN, else None."""
    if x is None:
        return None
    try:
        val = float(x)
    except Exception:
        return None
    if val != val:  # NaN check
        return None
    return round(val, 2)


def _mean_safe(values):
    vals = []
    for v in values:
        try:
            vals.append(float(v))
        except Exception:
            continue
    return mean(vals) if vals else None


def compute_kpis_from_json(j: dict) -> dict:
    """
    Compute KPIs from the dq_metrics JSON produced by the cleaning script.

    Focus is on completeness per modality:
      - completeness_gaze_pct:     mean completeness of gaze-related columns
      - completeness_input_pct:    mean completeness of input (KBM) columns
      - completeness_emotion_pct:  mean completeness of emotion-related columns
      - completeness_pct:          overall mean completeness across GAZE+INPUT+EMOTION
      - overall_quality_pct:       taken directly from JSON if present
      - uniqueness_pct / integrity_pct: structural indicators based on duplicate rows
    """

    after = j.get("after", j)  # fallback if 'after' not present

    rows = after.get("rows")
    duplicate_rows = after.get("duplicate_rows", 0)
    comp_by_col = after.get("completeness_pct_by_column", {}) or {}
    overall_quality = after.get("overall_quality_score_pct")

    # Restrict completeness to actual data columns (exclude metadata),
    # and only consider columns that are part of any modality group.
    all_modality_cols = GAZE_COLS | INPUT_COLS | EMOTION_COLS
    usable_cols = {
        c: v
        for c, v in comp_by_col.items()
        if c in all_modality_cols
    }

    # --- Overall completeness across all modalities ---
    completeness_overall = _mean_safe(usable_cols.values()) if usable_cols else None

    # --- Modality-specific completeness ---
    completeness_gaze = _mean_safe(
        usable_cols[c]
        for c in usable_cols.keys()
        if c in GAZE_COLS
    ) if usable_cols else None

    completeness_input = _mean_safe(
        usable_cols[c]
        for c in usable_cols.keys()
        if c in INPUT_COLS
    ) if usable_cols else None

    completeness_emotion = _mean_safe(
        usable_cols[c]
        for c in usable_cols.keys()
        if c in EMOTION_COLS
    ) if usable_cols else None

    # --- Uniqueness based on duplicate_rows ---
    uniqueness = None
    if isinstance(rows, int) and rows > 0:
        uniqueness = 100.0 * (1.0 - (float(duplicate_rows) / float(rows)))
        if uniqueness < 0:
            uniqueness = 0.0

    # --- Integrity as simple structural proxy (duplicates present or not) ---
    if duplicate_rows == 0:
        integrity = 100.0
    else:
        if isinstance(rows, int) and rows > 0:
            integrity = 100.0 * (1.0 - (float(duplicate_rows) / float(rows)))
            if integrity < 0:
                integrity = 0.0
        else:
            integrity = None

    return {
        "rows":                      rows,
        "completeness_pct":          pct(completeness_overall),
        "completeness_gaze_pct":     pct(completeness_gaze),
        "completeness_input_pct":    pct(completeness_input),
        "completeness_emotion_pct":  pct(completeness_emotion),
        "uniqueness_pct":            pct(uniqueness),
        "integrity_pct":             pct(integrity),
        "overall_quality_pct":       pct(overall_quality),
    }


def add_table(doc: Document, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"  # visible borders
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = "" if val is None else str(val)
    return t


# --------------------------
# Report builder
# --------------------------

def build_report(metrics_dir: Path, out_docx: Path):
    metrics_dir = Path(metrics_dir)
    out_docx = Path(out_docx)
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    files = sorted(metrics_dir.glob("*_dq_metrics.json"))
    if not files:
        raise SystemExit(f"No dq_metrics JSON files found in: {metrics_dir}")

    rows = []
    skipped = []
    for fp in files:
        match_prefix, player, pc = parse_ids_from_filename(fp)
        if not match_prefix or not player:
            skipped.append(fp.name)
            continue
        try:
            j = json.loads(fp.read_text(encoding="utf-8"))
        except Exception:
            skipped.append(fp.name)
            continue

        kpis = compute_kpis_from_json(j)

        # Build a combined player key so Pc1_P004 and Pc2_P004 are treated as different players
        if pc:
            player_id = f"Pc{pc}_{player}"
            pc_label = f"Pc{pc}"
        else:
            player_id = player           # old naming convention without PC info
            pc_label = None              # unknown / not encoded

        rows.append({
            "match_prefix": match_prefix,
            "player": player,        # raw player code (P004)
            "pc": pc,                # raw PC id (e.g. "1"), or None
            "pc_label": pc_label,    # formatted PC name, e.g. "Pc1"
            "player_id": player_id,  # combined key, e.g. "Pc1_P004"
            **kpis,
        })

    if not rows:
        raise SystemExit("No parsable files. Check filenames and JSON format.")

    df = pd.DataFrame(rows)

    # Aggregate per combined player (PC + player ID), per match, and per PC
    by_player = (
        df.groupby("player_id")
          .mean(numeric_only=True)
          .reset_index()
          .sort_values("player_id")
    )
    by_match = (
        df.groupby("match_prefix")
          .mean(numeric_only=True)
          .reset_index()
          .sort_values("match_prefix")
    )
    # PC-level aggregation (PC label, e.g. "Pc1", "Pc2").
    # Groupby ignores NaN by default, so old files without PC just drop out here.
    by_pc = (
        df.dropna(subset=["pc_label"])
          .groupby("pc_label")
          .mean(numeric_only=True)
          .reset_index()
          .sort_values("pc_label")
    )

    # Save helper CSVs
    report_dir = out_docx.parent
    by_player.to_csv(report_dir / "summary_players.csv", index=False)
    by_match.to_csv(report_dir / "summary_matches.csv", index=False)
    by_pc.to_csv(report_dir / "summary_pcs.csv", index=False)

    # Overall averages across all match–player files
    overall = df.mean(numeric_only=True).to_dict()

    # --------------------------
    # Build Word doc
    # --------------------------
    doc = Document()
    doc.add_heading("Data Quality Report – Breda Guardians ComicCon Dataset", level=1)

    # 1. Introduction (revised)
    doc.add_heading("1. Introduction", level=2)
    doc.add_paragraph(
        "This report provides an overview of data quality for the Breda Guardians ComicCon datasets, "
        "with a focus on recordings collected during VALORANT matches. The League of Legends data is "
        "excluded from this analysis because preliminary tests showed that the PCs could not reliably "
        "run both the game and the full Play-O-Meter toolkit simultaneously, resulting in incomplete "
        "or unusable recordings."
    )
    doc.add_paragraph(
        "For the Valorant matches, this report summarizes the data quality checks generated per match–player "
        "recording and aggregates them across players, matches, and PCs. The purpose is to determine the "
        "reliability of three core data streams—gaze, input, and emotion—and to identify where the toolkit "
        "performed well and where technical failures occurred."
    )

    # 2. Methodology (revised)
    doc.add_heading("2. Methodology", level=2)
    doc.add_paragraph(
        "All files matching the pattern '*_dq_metrics.json' are loaded from the selected directory. "
        "The match prefix (e.g., '6th_game'), player ID (e.g., 'P004'), and when present, the PC identifier "
        "(e.g., 'Pc3') is extracted from the filename. This allows differentiation between players who used "
        "different PCs (for example 'Pc1_P004' and 'Pc2_P004'), since these represent separate individuals or "
        "sessions."
    )
    doc.add_paragraph(
        "For each cleaned dataset, completeness is calculated per column and grouped into three modality "
        "categories: gaze (left/right gaze vectors and normalized screen coordinates), input (keyboard/mouse "
        "events and cursor positions), and emotion (emotion label and confidence values produced by the emotion "
        "recognition model). This modality-based approach highlights which parts of the pipeline performed well "
        "and which failed. Completeness is then aggregated per player, per match, and per PC."
    )
    doc.add_paragraph(
        "During data collection, several external factors affected data quality. Emotion completeness was the "
        "most problematic across sessions, often due to the camera freezing, the emotion model losing the face, "
        "or OBS video not saving during recording drops. Additionally, some players sat extremely close to their "
        "monitors, which significantly reduced the effectiveness of the eye-tracking system and contributed to "
        "lower gaze completeness in those sessions."
    )

    # 3. Results & Analysis
    doc.add_heading("3. Results & Analysis", level=2)
    doc.add_paragraph("Overall averages across all match–player files:")

    headers = ["KPI", "Average (%)"]
    rows_overall = [
        ["Completeness – Gaze",      pct(overall.get("completeness_gaze_pct"))],
        ["Completeness – Input",     pct(overall.get("completeness_input_pct"))],
        ["Completeness – Emotion",   pct(overall.get("completeness_emotion_pct"))],
        ["Completeness – Overall",   pct(overall.get("completeness_pct"))],
        ["Uniqueness",               pct(overall.get("uniqueness_pct"))],
        ["Integrity",                pct(overall.get("integrity_pct"))],
        ["Overall Quality",          pct(overall.get("overall_quality_pct"))],
    ]
    add_table(doc, headers, rows_overall)

    # Player-level
    doc.add_paragraph("\nPlayer-level summary (modality completeness):")
    headers_p = [
        "Player (PC+ID)",
        "Completeness – Gaze",
        "Completeness – Input",
        "Completeness – Emotion",
        "Completeness – Overall",
        "Overall Quality",
    ]
    rows_p = []
    for _, r in by_player.iterrows():
        rows_p.append([
            r["player_id"],
            pct(r.get("completeness_gaze_pct")),
            pct(r.get("completeness_input_pct")),
            pct(r.get("completeness_emotion_pct")),
            pct(r.get("completeness_pct")),
            pct(r.get("overall_quality_pct")),
        ])
    add_table(doc, headers_p, rows_p)

    # Match-level
    doc.add_paragraph("\nMatch-level summary (modality completeness):")
    headers_m = [
        "Match (Prefix)",
        "Completeness – Gaze",
        "Completeness – Input",
        "Completeness – Emotion",
        "Completeness – Overall",
        "Overall Quality",
    ]
    rows_m = []
    for _, r in by_match.iterrows():
        rows_m.append([
            r["match_prefix"],
            pct(r.get("completeness_gaze_pct")),
            pct(r.get("completeness_input_pct")),
            pct(r.get("completeness_emotion_pct")),
            pct(r.get("completeness_pct")),
            pct(r.get("overall_quality_pct")),
        ])
    add_table(doc, headers_m, rows_m)

    # PC-level
    if not by_pc.empty:
        doc.add_paragraph("\nPC-level summary (modality completeness):")
        headers_pc = [
            "PC",
            "Completeness – Gaze",
            "Completeness – Input",
            "Completeness – Emotion",
            "Completeness – Overall",
            "Overall Quality",
        ]
        rows_pc = []
        for _, r in by_pc.iterrows():
            rows_pc.append([
                r["pc_label"],
                pct(r.get("completeness_gaze_pct")),
                pct(r.get("completeness_input_pct")),
                pct(r.get("completeness_emotion_pct")),
                pct(r.get("completeness_pct")),
                pct(r.get("overall_quality_pct")),
            ])
        add_table(doc, headers_pc, rows_pc)

    # 4. Discussion (revised)
    doc.add_heading("4. Discussion", level=2)
    doc.add_paragraph(
        "Separating completeness into gaze, input, and emotion modalities reveals clear differences in subsystem "
        "performance. Emotion data shows the lowest completeness across the dataset. The primary causes identified "
        "during ComicCon were camera freezing, the OBS video capture failing to save, and cases where the player's "
        "face was not consistently visible (for example when players leaned extremely close to the monitor). These "
        "hardware and environmental issues made the emotion recognition subsystem significantly less reliable than "
        "the gaze or input trackers."
    )
    doc.add_paragraph(
        "Gaze data was generally more stable, but certain sessions exhibited reduced performance when players sat "
        "unusually close to the monitor. Under such conditions, the eye tracker could not consistently detect both "
        "eyes or maintain a stable gaze vector, leading to sparse or invalid gaze samples. Input data (mouse and "
        "keyboard events) was the most robust of the three modalities, with high completeness across nearly all "
        "recordings, confirming that the KBM subsystem was able to operate effectively even when other components "
        "failed."
    )
    doc.add_paragraph(
        "When comparing PCs, some machines consistently showed lower completeness scores, especially emotion and "
        "sometimes gaze. This suggests localized configuration or hardware issues, which helps explain why the League "
        "of Legends matches—run under heavier system load—were unusable. Only the Valorant sessions provided reliable "
        "multi-modality data for systematic analysis."
    )

    # 5. Recommendations (revised)
    doc.add_heading("5. Recommendations", level=2)
    doc.add_paragraph(
        "Based on the findings from the Valorant dataset, the following steps are recommended:\n"
        "• Investigate PCs with systematically low gaze or emotion completeness to identify hardware limitations, "
        "faulty cameras, or misconfigured eye-tracking hardware.\n"
        "• Ensure that OBS and the camera system are monitored during recording to prevent silent failures where "
        "video files are not saved.\n"
        "• Encourage players to maintain an appropriate distance from the monitor to improve the stability of gaze "
        "and emotion tracking.\n"
        "• Re-evaluate hardware requirements before running multi-stream recordings with demanding games like League "
        "of Legends, as previous attempts showed the PCs were not powerful enough to run both the game and the "
        "toolkit reliably.\n"
        "• Continue tracking modality-level completeness over future sessions to monitor improvements and detect "
        "regressions early.\n"
    )

    # 6. Conclusion (revised)
    doc.add_heading("6. Conclusion", level=2)
    doc.add_paragraph(
        "The Valorant dataset provides a workable foundation for AI coaching and analytic modeling, but quality varies "
        "significantly by modality. Input data remains highly reliable; gaze data is usable with expected variability, "
        "and emotion data is the least consistent due to camera and recording issues. By incorporating PC-level and "
        "modality-level quality metrics, this report highlights where the Play-O-Meter toolkit performed well and "
        "where technical challenges occurred. These insights help guide future improvements in data capture procedures, "
        "hardware selection, and operational protocols, especially for events where hardware stress and environmental "
        "conditions vary widely."
    )

    doc.add_paragraph(
        f"\nGenerated on {datetime.now().strftime('%d %B %Y at %H:%M')}."
    )

    doc.save(out_docx)
    print(f"Saved report → {out_docx}")
    print(
        f"Wrote CSVs → {report_dir / 'summary_players.csv'} , "
        f"{report_dir / 'summary_matches.csv'} , "
        f"{report_dir / 'summary_pcs.csv'}"
    )


# --------------------------
# CLI
# --------------------------

if __name__ == "__main__":
    p = argparse.ArgumentParser(
        description="Build Data Quality Report from dq_metrics JSON files."
    )
    p.add_argument(
        "--metrics-dir",
        required=True,
        help="Folder with *_dq_metrics.json files.",
    )
    p.add_argument(
        "--out",
        default="reports/ComicCon_Data_Quality_Report_BUas_Final.docx",
        help="Output .docx path.",
    )
    args = p.parse_args()
    build_report(Path(args.metrics_dir), Path(args.out))